from docx import Document
from random import randint
from datetime import date
from tinydb import TinyDB, Query
import json

from docx_image_replacer import ReplaceImage

# класс принимает данные json и генерирует файл docx на основе полученных данных
# в ответе отправляет путь до файла, название файлу и путь для создания папки на ядиске
class CreateDocx:
    def create(self, data):
        # print(data)
        # передаем в функцию выбранную смену и скидки. Обратно получаем словарь из обработанных цен
        prices = self.__mathPrices(data['shift'], data['sales_list'])
        print('Debug: Получили цены')
        # словарь, в котором перечислены вордовские переменные и, на что их заменить
        REPLACEMENTS = {
            'shift': data['shift'],
            'raw_price': prices['raw_price'],
            'result_price': prices['result_price'],
            'early_sale': prices['early_sale'],
            'sale_100': prices['sale_100'],
            'sale_2': prices['sale_2'],
            'sale_3': prices['sale_3'],
            'sale_4': prices['sale_4'],
            'frnd_sale': prices['frnd_sale'],
            'parent_fio': data['parent_fio'],
            'parent_passport': data['parent_passport'],
            'parent_reg': data['parent_reg'],
            'passport_date': data['passport_date'],
            'parent_address': data['parent_address'],
            'parent_phone': data['parent_phone'],
            'parent_email': data['parent_email'],
            'dop_fio': data['dop_parent_fio'],
            'dop_phone': data['dop_parent_phone'],
            'child_fio': data['child_fio'],
            'child_born': data['child_born'],
            'child_age': data['child_age'],
            'zabol_serd': data['zabol_serd'],
            'zabol_opor': data['zabol_opor'],
            'zabol_dyhat': data['zabol_dyhat'],
            'zabol_psyh': data['zabol_psyh'],
            'zabol_allerg': data['zabol_allerg'],
            'zabol_vnuo': data['zabol_vnuo'],
            'zabol_operat': data['zabol_operat'],
            'zabol_trawm': data['zabol_trawm']
        }
        CONTRACT_ID = str(randint(1000,9999)) # № договора
        TODAY = str(date.today()).split('-') # сегодняшняя дата
        USER_SIGNATURE = data['signature'] # подпись клиента в формате base64


        # открываем шаблон договора
        document = Document('docs/dogovor_template.docx')

        # замена переменных непосредственно в тексте (параграфы)
        for paragraph in document.paragraphs:
            inline = paragraph.runs # эта штука нужна, чтобы сохранить стили текста
            for i in range(len(inline)): # эта штука нужна, чтобы сохранить стили текста
                # print(paragraph.text)
                # print(inline[i].text)
                if 'contractid' in inline[i].text:
                    inline[i].text = inline[i].text.replace('contractid', CONTRACT_ID)
                if 'contract_date' in inline[i].text:
                    inline[i].text = inline[i].text.replace('contract_date', f'{TODAY[2]}.{TODAY[1]}.{TODAY[0]}')
                if 'parent_fio' in inline[i].text:
                    inline[i].text = inline[i].text.replace('parent_fio', REPLACEMENTS['parent_fio'])
                if 'parent_phone' in inline[i].text:
                    inline[i].text = inline[i].text.replace('parent_phone', REPLACEMENTS['parent_phone'])
                if 'child_fio' in inline[i].text:
                    inline[i].text = inline[i].text.replace('child_fio', REPLACEMENTS['child_fio'])

        # замена переменных внутри таблиц
        # проходимся по каждой таблице, каждой строке и каждой ячейке
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:

                    # как оказалось, у каждой ячейки есть т.н. параграфы
                    # print(cell.paragraphs)
                    # делим ячейки по параграфам
                    for paragraph in cell.paragraphs:
                        inline = paragraph.runs  # эта штука нужна, чтобы сохранить стили текста
                        for i in range(len(inline)):  # эта штука нужна, чтобы сохранить стили текста (обходим каждое слово параграфа)

                            # print(inline[i].text)
                            # print(cell.text)
                            # проходимся по каждому ключу в словаре REPLACEMENTS
                            for key in REPLACEMENTS.keys():

                                # если нашли блок с заболвеваниями Да/нет, то смотрим, есть ли в словаре замен что-то для вставки
                                # если есть, значит там указано заболевание. Ставим "Да"
                                # если пусто, значит заболевания не указано. Ставим "Нет"
                                if key + '_is' in inline[i].text:
                                    if REPLACEMENTS[key] != '':
                                        inline[i].text = inline[i].text.replace(key + '_is', 'Да')
                                    else:
                                        inline[i].text = inline[i].text.replace(key + '_is', 'Нет')

                                # если замечаем в ячейке переменную, заменяем её на текст из пришедшего сюда json
                                if key in inline[i].text:
                                    inline[i].text = inline[i].text.replace(key, str(REPLACEMENTS[key]))

        # генерируем новое имя для файла на основе рандомного числа и ФИО родителя
        # затем сохраняем файл в папке docs
        # возвращаем путь до файла
        doc_path = f'docs/{CONTRACT_ID}_{data["parent_fio"]}.docx'
        print('DEBUG: Создали путь файла')
        document.save(doc_path)
        print('DEBUG: Сохранили документ')

        # запускаем процесс замены шаблонной картинки подписи на подпись клиента
        # внутрь передаем png в формате base64 и путь к файлу, в котором нужно заменить картинку
        ReplaceImage(USER_SIGNATURE, doc_path)
        print('DEBUG: Заменили картинку')

        # путь для ядиска: "docs/1234_Иванов Иван Иванович (11.04.2022)/1234_Иванов Иван Иванович.docx"
        return doc_path, f'{CONTRACT_ID}_{data["parent_fio"].rstrip()}.docx',  f'docs/{CONTRACT_ID}_{data["parent_fio"]} ({TODAY[2]}.{TODAY[1]}.{TODAY[0]})'

    # функция принимает выбранную юзером смену и выбранные скидки
    # достает из бд все цены, считает результат и возвращает словарь из всех цен
    def __mathPrices(self, shift, sales_list):
        print('Debug: Зашли в функцию получения цен')
        # достаем цену на выбранную пользователем смену
        # ищем в бд смену с именем, пришедшем в запросе
        # проверяем, активна ли смена, не прошел ли срок вступления
        # берём цену этой смена -- получается сырая цена, из которой будут вычитаться скидки

        db = TinyDB('db/shift.json')
        user_shift = db.get(Query()['name'] == shift)

        if user_shift['isActive'] == False:
            return '[Ошибка] Данная смена закрыта'
        # проверка, если смена уже началась, -- отказать

        raw_price = user_shift['price']

        # достаем из бд скидки, выбранные пользователем в чекбоксах
        # на сервер выбранные чекбоксы приходят в виде списка с названиями скидок
        # ищем эти скидки в бд
        # проходимся по всем скидкам в бд и всем чекбоксам пользователя. Если нашли совпадение, то записываем скидку в нужную переменную
        # оставшиеся скидки так и останутся в переменной 0 руб
        early_sale = sale_100 = sale_2 = sale_3 = sale_4 = frnd_sale = 0

        db = TinyDB('db/sale.json')
        for sale in db.all():
            for i in json.loads(sales_list).values():
                if i in sale['name'] and sale['isActive'] == True:  # если нашли скидку из чекбокса, и она активна
                    # записываем скидки в переменные для вставки в договор
                    if sale['name'] == 'Оплачу 100% при заключении договора':
                        sale_100 = sale['sale']
                    elif sale['name'] == 'Поедут двое друзей':
                        frnd_sale = sale['sale']
                    elif sale['name'] == 'Ребенок участвует во второй раз':
                        sale_2 = sale['sale']
                    elif sale['name'] == 'Ребенок участвует в третий раз':
                        sale_3 = sale['sale']
                    elif sale['name'] == 'Ребенок участвует в четвертый раз':
                        sale_4 = sale['sale']
                    if sale['name'] == 'Уже есть страховка':
                        raw_price -= sale['sale']

        # получаем итоговую цену за смену
        # вычитаем из сырой цены все скидки
        result_price = raw_price - (early_sale + sale_100 + sale_2 + sale_3 + sale_4 + frnd_sale)

        return {
            'raw_price': raw_price,
            'result_price': result_price,
            'early_sale': early_sale,
            'sale_100': sale_100,
            'sale_2': sale_2,
            'sale_3': sale_3,
            'sale_4': sale_4,
            'frnd_sale': frnd_sale
        }